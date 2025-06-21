import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def is_hindi(text):
    # Checks for Devanagari characters (Hindi Unicode block)
    return bool(re.search(r'[\u0900-\u097F]', text))

def writeContentToDoc(file_path, content):
    """
    Writes or appends content to a .docx file.
    Supports both English and Hindi content with correct fonts.
    
    :param file_path: Path to the .docx file.
    :param content: String or list of strings to write to the file.
    """
    # Create folder path if it doesn't exist
    folder = os.path.dirname(file_path)
    if folder and not os.path.exists(folder):
        os.makedirs(folder, exist_ok=True)

    # Load or create the document
    if os.path.exists(file_path):
        doc = Document(file_path)
        doc.add_page_break()
    else:
        doc = Document()

    # Normalize content to list
    if isinstance(content, str):
        content = [content]

    for paragraph_text in content:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_text)

        if is_hindi(paragraph_text):
            run.font.name = 'Mangal'  # Use a font that supports Hindi
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Mangal')
        else:
            run.font.name = 'Calibri'  # Default English font

        run.font.size = Pt(14)

    doc.save(file_path)
    print(f"Content written/appended to: {file_path}")
